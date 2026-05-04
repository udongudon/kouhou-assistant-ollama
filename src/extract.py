"""議案書ファイル(PDF / Word / テキスト)からプレーンテキストを抽出する。"""
from __future__ import annotations

import io
from typing import IO


def extract_from_pdf(file_obj: IO[bytes]) -> str:
    from pypdf import PdfReader

    reader = PdfReader(file_obj)
    pages = []
    for page in reader.pages:
        text = page.extract_text() or ""
        pages.append(text)
    return "\n\n".join(pages).strip()


def extract_from_docx(file_obj: IO[bytes]) -> str:
    from docx import Document

    doc = Document(file_obj)
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()


def extract(file_obj: IO[bytes], filename: str) -> str:
    """ファイル拡張子で抽出方法を切り替える。"""
    name = filename.lower()
    if name.endswith(".pdf"):
        return extract_from_pdf(file_obj)
    if name.endswith(".docx"):
        return extract_from_docx(file_obj)
    if name.endswith((".txt", ".md")):
        data = file_obj.read()
        if isinstance(data, bytes):
            for encoding in ("utf-8", "cp932", "shift_jis"):
                try:
                    return data.decode(encoding).strip()
                except UnicodeDecodeError:
                    continue
            return data.decode("utf-8", errors="replace").strip()
        return str(data).strip()
    raise ValueError(
        f"未対応のファイル形式です: {filename}. PDF / DOCX / TXT をアップロードしてください。"
    )


def truncate_for_prompt(text: str, max_chars: int = 30000) -> tuple[str, bool]:
    """プロンプトに渡す前に長すぎる議案書を切り詰める。

    切り詰めた場合は (text, True) を返す。
    """
    if len(text) <= max_chars:
        return text, False
    return text[:max_chars], True
